from __future__ import annotations

import csv
import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
TABLE_DIR = ROOT / "03_results_tables"
LOG_DIR = ROOT / "05_logs"
MS_DIR = ROOT / "06_manuscript_support"
SUB_DIR = ROOT / "07_submission_package"

for p in [
    TABLE_DIR,
    LOG_DIR,
    MS_DIR,
    SUB_DIR / "manuscript",
    SUB_DIR / "cover_letter",
    SUB_DIR / "data_code_availability",
    SUB_DIR / "reporting_checklists",
]:
    p.mkdir(parents=True, exist_ok=True)


SCAN_FILES = [
    MS_DIR / "phase10A_fix_manuscript_reference_enhanced.txt",
    MS_DIR / "phase10E_cover_letter_Human_Genomics.txt",
    MS_DIR / "phase10E_graphical_abstract_design_brief.txt",
    MS_DIR / "phase10E_data_availability_statement.txt",
    MS_DIR / "phase10E_code_availability_statement.txt",
    MS_DIR / "phase10E_reproducibility_README.txt",
    MS_DIR / "phase10E_declarations_template.txt",
    MS_DIR / "phase10D_final_main_figure_legends.txt",
    LOG_DIR / "phase10E_key_result_checklist.txt",
    LOG_DIR / "phase10D_key_result_checklist.txt",
]

PLACEHOLDER_RE = re.compile(
    r"\[[^\]]*(?:AUTHOR CONFIRMATION NEEDED|REPOSITORY NEEDED|FUNDING NEEDED|ORCID NEEDED|ETHICS CONFIRMATION NEEDED|NEEDED|CONFIRMATION)[^\]]*\]",
    re.IGNORECASE,
)


def infer_section(text: str, pos: int) -> str:
    before = text[:pos].splitlines()
    for line in reversed(before[-80:]):
        stripped = line.strip()
        if stripped.startswith("#"):
            return stripped.lstrip("#").strip()
        if stripped and stripped.lower() in {
            "availability of data and materials",
            "code availability",
            "funding",
            "author contributions",
            "competing interests",
            "ethics approval and consent to participate",
        }:
            return stripped
    return "general"


def placeholder_inventory() -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for file in SCAN_FILES:
        if not file.exists():
            continue
        text = file.read_text(encoding="utf-8", errors="ignore")
        for match in PLACEHOLDER_RE.finditer(text):
            ph = match.group(0)
            lower = ph.lower()
            required = "yes"
            priority = "must"
            if "optional" in lower or "suggested reviewers" in lower or "opposed reviewers" in lower:
                required = "no"
                priority = "optional"
            elif "orcid" in lower:
                required = "no"
                priority = "recommended"
            if "repository" in lower or "url" in lower or "doi" in lower:
                suggestion = "Provide final repository name, URL, DOI or stable identifier; do not invent one before upload."
            elif "funding" in lower or "grant" in lower:
                suggestion = "Provide funding source, grant number and recipient; or confirm no specific funding."
            elif "competing" in lower:
                suggestion = "Confirm no competing interests or provide the exact disclosure."
            elif "ethics" in lower or "institutional" in lower:
                suggestion = "Confirm institutional wording for public de-identified data and no new human/animal experiments."
            elif "author" in lower or "corresponding" in lower:
                suggestion = "Provide final author/corresponding-author information and author approval."
            else:
                suggestion = "Replace with final author-confirmed submission information."
            rows.append(
                {
                    "file": str(file.relative_to(ROOT)),
                    "section": infer_section(text, match.start()),
                    "placeholder_text": ph,
                    "required_for_submission": required,
                    "suggested_user_input": suggestion,
                    "priority": priority,
                }
            )
    out = TABLE_DIR / "phase10F_placeholder_inventory.csv"
    with out.open("w", newline="", encoding="utf-8-sig") as fh:
        writer = csv.DictWriter(
            fh,
            fieldnames=[
                "file",
                "section",
                "placeholder_text",
                "required_for_submission",
                "suggested_user_input",
                "priority",
            ],
        )
        writer.writeheader()
        writer.writerows(rows)
    return rows


def author_template_txt() -> str:
    return """
# Phase 10F author information template

Please fill one block per author. Do not leave required fields blank before submission.

## Author 1
Full author name:
Affiliation:
Email:
ORCID:
Author order:
Equal contribution yes/no:
Corresponding author yes/no:

Contribution categories, mark yes/no or describe role:
- Conceptualization:
- Data curation:
- Formal analysis:
- Methodology:
- Visualization:
- Writing - original draft:
- Writing - review & editing:
- Supervision:
- Funding acquisition:

Funding source:
Conflict of interest confirmation:

## Author 2
Full author name:
Affiliation:
Email:
ORCID:
Author order:
Equal contribution yes/no:
Corresponding author yes/no:

Contribution categories, mark yes/no or describe role:
- Conceptualization:
- Data curation:
- Formal analysis:
- Methodology:
- Visualization:
- Writing - original draft:
- Writing - review & editing:
- Supervision:
- Funding acquisition:

Funding source:
Conflict of interest confirmation:

## Notes for the corresponding author
- Confirm final author order before uploading the manuscript.
- Confirm whether any authors contributed equally.
- Confirm all CRediT roles.
- Confirm competing interests and funding statements exactly.
- Provide ORCID IDs if required by the submission system.
"""


def author_template_xlsx() -> Path:
    headers = [
        "Full author name",
        "Affiliation",
        "Email",
        "ORCID",
        "Author order",
        "Equal contribution yes/no",
        "Corresponding author yes/no",
        "Conceptualization",
        "Data curation",
        "Formal analysis",
        "Methodology",
        "Visualization",
        "Writing - original draft",
        "Writing - review & editing",
        "Supervision",
        "Funding acquisition",
        "Funding source",
        "Conflict of interest confirmation",
    ]
    wb = Workbook()
    ws = wb.active
    ws.title = "Author information"
    ws.append(headers)
    for _ in range(12):
        ws.append([""] * len(headers))
    fill = PatternFill("solid", fgColor="D9EAF7")
    thin = Side(style="thin", color="B7B7B7")
    border = Border(top=thin, bottom=thin, left=thin, right=thin)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
    for row in ws.iter_rows(min_row=2, max_row=13, max_col=len(headers)):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border
    widths = [22, 32, 26, 20, 12, 20, 22] + [20] * 9 + [28, 32]
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    notes = wb.create_sheet("Instructions")
    notes["A1"] = "Instructions"
    notes["A1"].font = Font(bold=True, size=14)
    instructions = [
        "Fill one row per author.",
        "Use yes/no or concise role descriptions for CRediT categories.",
        "Do not invent ORCID, funding or competing-interest details.",
        "Confirm equal-contribution and corresponding-author status before submission.",
        "If no funding exists, write: This research received no specific grant from any funding agency in the public, commercial, or not-for-profit sectors.",
        "If no competing interests exist, write: The authors declare no competing interests.",
    ]
    for i, item in enumerate(instructions, 3):
        notes[f"A{i}"] = item
        notes[f"A{i}"].alignment = Alignment(wrap_text=True)
    notes.column_dimensions["A"].width = 120
    out = MS_DIR / "phase10F_author_information_template.xlsx"
    wb.save(out)
    return out


def write(path: Path, text: str) -> Path:
    path.write_text(text.strip() + "\n", encoding="utf-8")
    return path


def declarations_final() -> str:
    return """
# Ethics approval and consent to participate

This study used publicly available, de-identified datasets and involved no new human participants or animal experiments. [ETHICS CONFIRMATION NEEDED: confirm whether the authors' institution requires an ethics-exemption statement or additional wording.]

Optional wording if institution confirms no additional review was required:
"Because all analyses used publicly available, de-identified datasets, no additional institutional ethics approval was required."

# Consent for publication

Not applicable.

# Competing interests

[AUTHOR CONFIRMATION NEEDED: choose one option before submission.]

Option A, if no competing interests exist:
"The authors declare no competing interests."

Option B, if competing interests exist:
"The authors declare the following competing interests: [AUTHOR CONFIRMATION NEEDED]."

# Funding

[FUNDING NEEDED: provide funding source, grant number and recipient, or confirm no specific funding.]

Optional wording if no specific funding was received:
"This research received no specific grant from any funding agency in the public, commercial, or not-for-profit sectors."

# Author contributions

Conceptualization: [AUTHOR CONFIRMATION NEEDED]
Data curation: [AUTHOR CONFIRMATION NEEDED]
Formal analysis: [AUTHOR CONFIRMATION NEEDED]
Methodology: [AUTHOR CONFIRMATION NEEDED]
Visualization: [AUTHOR CONFIRMATION NEEDED]
Writing - original draft: [AUTHOR CONFIRMATION NEEDED]
Writing - review & editing: [AUTHOR CONFIRMATION NEEDED]
Supervision: [AUTHOR CONFIRMATION NEEDED]
Funding acquisition: [AUTHOR CONFIRMATION NEEDED]

All authors read and approved the final manuscript. [AUTHOR CONFIRMATION NEEDED]

# Acknowledgements

[AUTHOR CONFIRMATION NEEDED: acknowledge colleagues, institutional support, language editing, data providers or computational resources as appropriate. If none, confirm whether this section should be omitted.]
"""


def cover_letter_final() -> str:
    src = MS_DIR / "phase10E_cover_letter_Human_Genomics.txt"
    text = src.read_text(encoding="utf-8") if src.exists() else ""
    replacements = {
        "[AUTHOR CONFIRMATION NEEDED: confirm final author approval.]": "[AUTHOR CONFIRMATION NEEDED: confirm all authors have approved this submission.]",
        "[AUTHOR CONFIRMATION NEEDED: confirm competing interests.]": "[AUTHOR CONFIRMATION NEEDED: confirm competing interests statement.]",
        "[AUTHOR CONFIRMATION NEEDED: Corresponding author name]": "[CORRESPONDING AUTHOR NAME]",
        "[AUTHOR CONFIRMATION NEEDED: Corresponding author affiliation]": "[CORRESPONDING AUTHOR AFFILIATION]",
        "[AUTHOR CONFIRMATION NEEDED: Corresponding author email]": "[CORRESPONDING AUTHOR EMAIL]",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def make_docx_from_text(text: str, out: Path, title: str) -> Path:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(16)
    for block in text.strip().split("\n\n"):
        if block.startswith("# "):
            p = doc.add_paragraph()
            rr = p.add_run(block[2:].strip())
            rr.bold = True
            rr.font.name = "Arial"
            rr.font.size = Pt(13)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            rr = p.add_run(block.replace("\n", " "))
            rr.font.name = "Arial"
            rr.font.size = Pt(11)
    doc.save(out)
    return out


def data_code_final() -> str:
    return """
# Final fillable Availability of data and materials

The study used publicly available, de-identified datasets and derived analysis outputs. TCGA-LAML transcriptomic and clinical data were obtained from public TCGA/GDC or UCSC Xena resources. GEO datasets GSE37642 and GSE12417 were used for formal external validation. GSE6891 and GSE14468 were audited but not used as formal validation cohorts because the locally available files lacked sufficient six-gene coverage and usable overall-survival endpoints. BeatAML expression and ex vivo drug-response data were used for pharmacogenomic association analyses. Processed single-cell tables from GSE116256 were used for processed single-cell table-based patient-level analyses. STRING was used for protein-protein interaction evidence.

Derived datasets supporting the conclusions of this article are included in the supplementary tables and will be deposited in [REPOSITORY NEEDED: repository name and URL] before publication. These derived files include PRFT scores, differential-expression and WGCNA outputs, candidate-gene lists, survival-validation summaries, clinical Cox outputs, machine-learning performance tables, PPI/network-prioritization tables, BeatAML ex vivo pharmacogenomic association summaries, processed single-cell patient-level summaries, bulk signature/pathway summaries, figure source audits and excluded-analysis audit tables.

No newly generated raw human sequencing data were produced in this study. No individual participant-level identifiable data are included in the manuscript or supplementary files.

# Final fillable Code availability

Code generated for this study will be deposited in [REPOSITORY NEEDED: GitHub/Zenodo/OSF/institutional repository URL] before publication. The archived version, license and permanent identifier are pending: [REPOSITORY NEEDED: DOI or stable URL]; [REPOSITORY NEEDED: software license]; [REPOSITORY NEEDED: release date].

# Repository options

Option 1, GitHub + Zenodo DOI:
- Upload scripts and README to GitHub.
- Create a tagged release.
- Archive the release through Zenodo and add the Zenodo DOI to the manuscript.

Option 2, OSF:
- Create an OSF project.
- Upload scripts, derived tables, README and figure source audit files.
- Use the OSF stable URL or DOI if minted.

Option 3, institutional repository:
- Deposit scripts, derived data and README in the institutional repository.
- Add the repository URL, DOI or handle supplied by the institution.

# Repository fields to fill before submission

Repository name:
Repository URL:
Repository DOI or stable identifier:
Software license:
Release date:
Version/tag:
Contact person for repository:
"""


def repository_upload_checklist() -> str:
    return """
# Repository upload file checklist

## Must upload before publication
1. Final manuscript source text or accepted manuscript version if allowed by journal policy.
2. All scripts needed to reproduce the reported analyses, organized by phase.
3. Derived result tables from 03_results_tables that support main and supplementary results.
4. Main figure source audit and figure-text consistency audit.
5. Supplementary table manifest and copied supplementary table files.
6. README describing project structure, data sources, software versions, reproduction steps and limitations.
7. Session information files and package-version logs.
8. License file for code reuse. [REPOSITORY NEEDED: choose license]

## Do not upload unless permissions allow
1. Restricted or controlled-access raw data.
2. Any identifiable participant-level information.
3. Temporary local files, absolute private paths or cache directories.
4. Placeholder UMAP figures or excluded AS result-like figures.

## Recommended repository organization
- README.md
- LICENSE
- scripts/
- results_tables/
- figure_source_audit/
- logs/
- manuscript_support/

## Final repository metadata to record
- Repository URL:
- DOI or stable identifier:
- Version/tag:
- Release date:
- License:
"""


def formatting_checklist() -> str:
    return """
# Phase 10F final Human Genomics formatting checklist

1. Title page: add final author names, affiliations, corresponding author, email and ORCID where required.
2. Abstract: confirm structured format and word count.
3. Keywords: confirm journal-compatible keywords.
4. Main text: ensure final DOCX contains line numbers and page numbers.
5. References: final list contains 61 references; complete manual metadata check for any remaining DOI/PMID exceptions.
6. Figure legends: use Phase 10D final legends and keep Figure 1-7 order.
7. Main Figures 1-7: visually inspect final PDF/PNG files before upload.
8. Supplementary Figures S1-S10: confirm final files or plan according to submission system.
9. Supplementary Tables S1-S15: confirm all files copied and names are submission-friendly.
10. Line numbers: add before submission.
11. Page numbers: add before submission.
12. Double spacing if required: check final Human Genomics submission instructions and apply if required.
13. Figure resolution: Phase 10D PNGs are 300 dpi; confirm no text is blurred after upload.
14. File naming: use clear names without internal phase wording if the submission system exposes file names to reviewers.
15. Declarations: fill ethics, competing interests, funding, author contributions and acknowledgements.
16. Data/code availability: replace repository URL/DOI/license placeholders after upload.
17. Ethics: confirm institutional wording for public de-identified datasets.
18. Cover letter: replace corresponding-author placeholders and confirm author approval.
19. Graphical abstract: review final BioRender/Illustrator artwork and ensure bounded-model language.
20. Repository link: verify all URLs resolve before submission.
"""


def readiness_summary_cn() -> str:
    return """
# Phase 10F 最终可投稿性判断

## 1. 现在是否可以立即投稿？
暂时不建议立即投稿。数据核心、主图、参考文献、cover letter、data/code availability 和 declarations 模板已经准备好，但仍缺少投稿系统必须的人工信息，包括作者、基金、利益冲突、伦理确认、仓库 URL/DOI/license 和最终图件视觉确认。

## 2. 还差哪些人工信息？
- 通讯作者姓名、单位和邮箱；
- 作者顺序、单位、ORCID 和是否共同一作/共同通讯；
- CRediT 作者贡献；
- 基金来源、课题号和受资助作者，或确认无特定基金；
- 利益冲突声明；
- 伦理/单位确认措辞；
- 数据和代码仓库 URL、DOI、license、版本号；
- Figure 1-7 最终视觉确认；
- 可选审稿人和回避审稿人。

## 3. 哪些文件已经可用？
- 参考文献增强版 manuscript；
- Figure 1-7 主图 PDF/PNG；
- Figure legends；
- Supplementary Tables S1-S15 manifest 和已复制表格；
- Graphical abstract brief 和绘图 prompt；
- Cover letter 可填写版；
- Data/code availability 可填写版；
- Declarations 可填写版；
- 作者信息 xlsx/txt 模板；
- 最终格式化 checklist。

## 4. 哪些文件必须由用户补全？
- phase10F_author_information_template.xlsx；
- phase10F_declarations_final_fillable.txt；
- phase10F_cover_letter_final_fillable.txt/docx；
- phase10F_data_code_availability_final_fillable.txt；
- repository upload checklist 中的 URL/DOI/license/release date；
- 最终 manuscript title page。

## 5. 哪些内容不能再改？
- 固定 6 基因模型：CLCN5, ITGB2, ARHGEF5, TRIM32, SAT1, ACOX2；
- AS 删除边界；
- processed single-cell table-based 分析边界；
- bulk signature-based inference 边界；
- BeatAML ex vivo pharmacogenomic association 边界；
- 不把 PPI 写成机制证明，不把 BeatAML 写成临床治疗推荐。

## 6. 下一步最短路径
先填写作者信息模板和 declarations；上传 derived tables/scripts 到 GitHub+Zenodo、OSF 或单位仓库；替换 Data/Code availability 中的仓库占位符；打开 Figure 1-7 做最终视觉确认；最后给 manuscript DOCX 添加 title page、line numbers、page numbers 和最终投稿格式。
"""


def checklist_text() -> str:
    return """
1. Placeholder inventory generated: yes
2. Author information template generated: yes
3. Declarations fillable version generated: yes
4. Cover letter fillable version generated: yes
5. Data/code availability fillable version generated: yes
6. Repository upload checklist generated: yes
7. Final formatting checklist generated: yes
8. Submission readiness summary generated: yes
9. Fixed six-gene model preserved: yes
10. AS exclusion preserved: yes
11. Processed single-cell boundary preserved: yes
12. Bulk signature-based inference boundary preserved: yes
13. BeatAML ex vivo boundary preserved: yes
14. Clinical recommendation avoided: yes
15. Can submit immediately: no
16. Must complete before immediate submission: corresponding author details; author order/contributions; funding; competing interests; ethics wording; repository URL/DOI/license; final figure visual approval; line/page numbering
17. Issues requiring human confirmation: author metadata, ORCID, funding, competing interests, ethics exemption wording, repository upload details, optional reviewer suggestions, final graphical abstract artwork
18. Overall conclusion: Phase 10F produced fillable final submission texts and templates; the package is near submission-ready but requires author-confirmed metadata before upload.
"""


def copy_to_submission(paths: list[Path]) -> None:
    for path in paths:
        if not path.exists():
            continue
        if "cover_letter" in path.name:
            dest_dir = SUB_DIR / "cover_letter"
        elif "data_code" in path.name or "repository" in path.name:
            dest_dir = SUB_DIR / "data_code_availability"
        elif "declarations" in path.name:
            dest_dir = SUB_DIR / "manuscript"
        elif "author_information" in path.name or "formatting_checklist" in path.name or "placeholder" in path.name or "readiness" in path.name:
            dest_dir = SUB_DIR / "reporting_checklists"
        else:
            dest_dir = SUB_DIR / "reporting_checklists"
        dest_dir.mkdir(parents=True, exist_ok=True)
        shutil.copy2(path, dest_dir / path.name)


def main() -> None:
    inventory = placeholder_inventory()
    author_txt = write(MS_DIR / "phase10F_author_information_template.txt", author_template_txt())
    author_xlsx = author_template_xlsx()
    declarations = write(MS_DIR / "phase10F_declarations_final_fillable.txt", declarations_final())
    cover_txt = write(MS_DIR / "phase10F_cover_letter_final_fillable.txt", cover_letter_final())
    cover_docx = make_docx_from_text(cover_letter_final(), MS_DIR / "phase10F_cover_letter_final_fillable.docx", "Cover Letter - Fillable Draft")
    data_code = write(MS_DIR / "phase10F_data_code_availability_final_fillable.txt", data_code_final())
    repo = write(MS_DIR / "phase10F_repository_upload_file_checklist.txt", repository_upload_checklist())
    fmt = write(MS_DIR / "phase10F_final_formatting_checklist.txt", formatting_checklist())
    summary = write(MS_DIR / "phase10F_submission_readiness_summary_CN.txt", readiness_summary_cn())
    phase_log = write(LOG_DIR / "phase10F_key_result_checklist.txt", checklist_text())
    copy_to_submission(
        [
            TABLE_DIR / "phase10F_placeholder_inventory.csv",
            author_txt,
            author_xlsx,
            declarations,
            cover_txt,
            cover_docx,
            data_code,
            repo,
            fmt,
            summary,
            phase_log,
        ]
    )


if __name__ == "__main__":
    main()
