from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
PACKAGE = ROOT / "Human_Genomics_PRFT_AML_submission_package"
FIGURE_PATH = PACKAGE / "figures" / "Figure2_PRFT_DEG_WGCNA.png"
REVIEW_DOCS = [
    PACKAGE / "Human_Genomics_PRFT_AML_review_copy_with_figures.docx",
    PACKAGE / "Human_Genomics_PRFT_AML_review_copy_with_figures_revised.docx",
]


def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)
    paragraph._p = paragraph._element = None


def add_paragraph_after(paragraph):
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def update_review_doc(docx_path: Path, figure_path: Path) -> bool:
    if not docx_path.exists():
        return False

    doc = Document(docx_path)
    anchor_idx = None
    figure_section_seen = False

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == "Embedded figures for review copy":
            figure_section_seen = True
            continue
        if figure_section_seen and text == "Figure 2":
            anchor_idx = idx
            break

    if anchor_idx is None:
        return False

    anchor = doc.paragraphs[anchor_idx]

    next_idx = anchor_idx + 1
    if next_idx < len(doc.paragraphs):
        next_para = doc.paragraphs[next_idx]
        has_drawing = any("drawing" in run._element.xml for run in next_para.runs)
        if next_para.text.strip() == "" or has_drawing:
            remove_paragraph(next_para)

    new_para = add_paragraph_after(anchor)
    new_para.alignment = 1
    run = new_para.add_run()
    run.add_picture(str(figure_path), width=Inches(7.2))

    doc.save(docx_path)
    return True


def main():
    if not FIGURE_PATH.exists():
        raise FileNotFoundError(f"Optimized Figure 2 not found: {FIGURE_PATH}")

    updated = []
    for docx_path in REVIEW_DOCS:
        if update_review_doc(docx_path, FIGURE_PATH):
            updated.append(str(docx_path))

    if not updated:
        raise RuntimeError("No review copy DOCX could be updated with the optimized Figure 2.")

    for item in updated:
        print(item)


if __name__ == "__main__":
    main()
