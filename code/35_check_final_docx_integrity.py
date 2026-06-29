from pathlib import Path

from docx import Document

base = Path(__file__).resolve().parents[1] / "Human_Genomics_PRFT_AML_submission_package_clean"
files = [
    base / "Human_Genomics_PRFT_AML_main_manuscript.docx",
    base / "Human_Genomics_PRFT_AML_review_copy_with_figures_final.docx",
    base / "cover_letter_draft.docx",
]

for f in files:
    doc = Document(str(f))
    rel_imgs = [r for r in doc.part.rels.values() if "image" in r.reltype]
    print(f.name)
    print("  exists:", f.exists(), "size_MB:", round(f.stat().st_size / 1024 / 1024, 3))
    print("  paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables), "embedded_images:", len(rel_imgs))
