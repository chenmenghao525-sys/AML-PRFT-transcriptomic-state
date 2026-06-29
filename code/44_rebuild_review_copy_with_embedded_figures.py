from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OFFICIAL = ROOT / "Human_Genomics_PRFT_AML_OFFICIAL_submission_package"
INTERNAL = ROOT / "Human_Genomics_PRFT_AML_INTERNAL_working_package"
OFFICIAL_ZIP = ROOT / "Human_Genomics_PRFT_AML_OFFICIAL_submission_package.zip"
INTERNAL_ZIP = ROOT / "Human_Genomics_PRFT_AML_INTERNAL_working_package.zip"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_image_block(doc: Document, title: str, image_path: Path, width_inches: float = 6.6) -> None:
    doc.add_page_break()
    add_heading(doc, title, level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    cap = doc.add_paragraph(title)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def rebuild_review_copy(root: Path) -> Path:
    main_doc = root / "01_Main_Manuscript" / "Human_Genomics_PRFT_AML_main_manuscript_final_clean.docx"
    review_doc = root / "01_Main_Manuscript" / "Human_Genomics_PRFT_AML_review_copy_with_figures_final_clean.docx"
    doc = Document(main_doc)

    add_heading(doc, "Embedded Figure Preview", level=1)
    doc.add_paragraph(
        "This review copy embeds figure previews after the clean manuscript text. "
        "The official submission package also provides separate high-resolution figure files."
    )

    graphical = root / "03_Graphical_Abstract" / "graphical_abstract_PRFT_AML.png"
    if graphical.exists():
        add_image_block(doc, "Graphical Abstract", graphical, width_inches=6.6)

    for i in range(1, 9):
        img = root / "04_Main_Figures" / f"Figure{i}_final.png"
        if i in {7, 8}:
            readable = root / "04_Main_Figures" / f"Figure{i}_final_readable.png"
            if readable.exists():
                img = readable
        if img.exists():
            add_image_block(doc, f"Figure {i}", img, width_inches=6.6)

    add_heading(doc, "Supplementary Figure Preview", level=1)
    for i in range(1, 5):
        img = root / "05_Supplementary_Figures" / f"Supplementary_Figure_S{i}.png"
        if img.exists():
            add_image_block(doc, f"Supplementary Figure S{i}", img, width_inches=6.6)

    doc.save(review_doc)
    return review_doc


def zip_dir(folder: Path, zip_path: Path) -> None:
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in folder.rglob("*"):
            if path.is_file():
                zf.write(path, path.relative_to(folder.parent))
    with zipfile.ZipFile(zip_path, "r") as zf:
        bad = zf.testzip()
    if bad is not None:
        raise RuntimeError(f"ZIP integrity failed at {bad}")


def main() -> None:
    official_review = rebuild_review_copy(OFFICIAL)
    internal_review = INTERNAL / "01_Main_Manuscript" / official_review.name
    internal_review.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(official_review, internal_review)
    zip_dir(OFFICIAL, OFFICIAL_ZIP)
    zip_dir(INTERNAL, INTERNAL_ZIP)
    print(official_review)
    print(internal_review)
    print(OFFICIAL_ZIP)
    print(INTERNAL_ZIP)


if __name__ == "__main__":
    main()
