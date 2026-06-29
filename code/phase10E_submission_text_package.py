from __future__ import annotations

import csv
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1]
TABLE_DIR = ROOT / "03_results_tables"
LOG_DIR = ROOT / "05_logs"
MS_DIR = ROOT / "06_manuscript_support"
SUB_DIR = ROOT / "07_submission_package"

for path in [
    TABLE_DIR,
    LOG_DIR,
    MS_DIR,
    SUB_DIR / "cover_letter",
    SUB_DIR / "data_code_availability",
    SUB_DIR / "reporting_checklists",
    SUB_DIR / "graphical_abstract",
]:
    path.mkdir(parents=True, exist_ok=True)


TITLE = (
    "A proteostasis-associated ferroptosis-tolerance signature defines adverse "
    "prognosis, myeloid immune suppression and ex vivo pharmacogenomic "
    "vulnerability in acute myeloid leukemia"
)

SIX_GENES = "CLCN5, ITGB2, ARHGEF5, TRIM32, SAT1 and ACOX2"


def write_text(name: str, text: str) -> Path:
    path = MS_DIR / name
    path.write_text(text.strip() + "\n", encoding="utf-8")
    return path


def copy_if_exists(src: Path, subfolder: str, dest_name: str | None = None) -> None:
    if src.exists():
        dest = SUB_DIR / subfolder / (dest_name or src.name)
        dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, dest)


def graphical_abstract_design_brief() -> str:
    return f"""
# Phase 10E graphical abstract design brief

## Purpose
Create a clean biomedical schematic for Human Genomics that summarizes the bounded PRFT-high AML state model. The graphic should communicate the study as a retrospective, public-data, computational and ex vivo pharmacogenomic analysis. It must not imply a proved molecular mechanism, clinical treatment recommendation or completed wet-lab validation.

## Overall layout
Use a left-to-right flow with five visual zones:

1. Input data layer: AML bulk transcriptomes and public datasets.
2. PRFT state definition: proteostasis-core score plus ferroptosis-tolerance score defines PRFT-high versus PRFT-low AML.
3. Fixed six-gene signature: {SIX_GENES}.
4. Evidence layers linked to PRFT-high AML:
   - adverse prognosis and external validation;
   - clinical independence and survival machine-learning stability;
   - monocyte-like/myeloid immune-suppressive transcriptional state;
   - JAK/STAT/PD-L1-related signatures;
   - oxidative stress/NRF2-related signatures;
   - ferroptosis-defense / SLC7A11-GPX4-GSH axis;
   - proteostasis / HSP90 / UPR axis;
   - BeatAML ex vivo pharmacogenomic associations.
5. Boundary box: retrospective public datasets; ex vivo association; no clinical treatment recommendation; no de novo single-cell reanalysis; no formal immune deconvolution; no AS module; experimental validation required.

## Visual style
Use a clean biomedical schematic, BioRender-like vector style, white background, restrained blue/teal/red-gold palette and simple cell/pathway icons. Avoid photorealistic patient images. Avoid dense pathway wiring. Use clear arrows labelled "associated with", "linked to" or "supports a bounded model". Do not use "causes", "activates", "proves", "reverses" or "validated mechanism".

## Recommended labels
- AML bulk transcriptomes
- PRFT state definition
- Proteostasis score
- Ferroptosis-tolerance score
- Fixed six-gene PRFT signature
- Adverse prognosis
- Myeloid/monocyte-like immune-suppressive transcriptional state
- JAK/STAT/PD-L1
- Oxidative stress/NRF2
- SLC7A11-GPX4-GSH
- Proteostasis/HSP90/UPR
- BeatAML ex vivo pharmacogenomic associations
- Experimental validation required

## Main visual boundary
The graphical abstract should read as an integrated state model. It should not look like a mechanistic pathway has been experimentally proven. The final right-side or bottom boundary box should be visible enough that reviewers can immediately see the study's limits.
"""


def graphical_abstract_prompt() -> str:
    return f"""
# English prompt for BioRender / AI-assisted vector drafting

Create a clean biomedical vector-style graphical abstract for a Human Genomics manuscript on acute myeloid leukemia. Use a left-to-right flow. On the left, show public AML bulk transcriptomic datasets feeding into a PRFT state definition: proteostasis-core score plus ferroptosis-tolerance score. In the center, show a fixed six-gene PRFT signature labelled "{SIX_GENES}". From this central PRFT-high AML state, draw restrained arrows labelled "associated with" or "linked to" toward four evidence modules: adverse prognosis and external validation; monocyte-like/myeloid immune-suppressive transcriptional state with JAK/STAT/PD-L1 and oxidative stress/NRF2 labels; ferroptosis-defense and proteostasis axes labelled SLC7A11-GPX4-GSH and HSP90/UPR; and BeatAML ex vivo pharmacogenomic associations. Use simple AML cell icons, transcriptome icons, network icons and drug plate icons. Use a white background, teal/blue/gray primary palette with modest red/gold accents for adverse prognosis and pharmacogenomics. Add a visible boundary box: "Retrospective public datasets; ex vivo association; experimental validation required; no clinical treatment recommendation." Do not draw real patient photographs. Do not draw an overly complex pathway. Do not use the words "proved mechanism", "causes", "activates" or "clinical recommendation".

# 中文绘图提示词

绘制一张用于 Human Genomics 投稿的 AML PRFT 图形摘要，风格为干净、BioRender-like、矢量医学示意图。整体采用从左到右的流程：左侧为公开 AML bulk transcriptome 数据输入；中间为 PRFT 状态定义，即 proteostasis-core score 加 ferroptosis-tolerance score；中心突出固定六基因 PRFT signature：{SIX_GENES}。从 PRFT-high AML state 向右侧分出四个证据模块，箭头文字使用 "associated with" 或 "linked to"：不良预后和外部验证；monocyte-like/myeloid immune-suppressive transcriptional state，并标注 JAK/STAT/PD-L1 和 oxidative stress/NRF2；ferroptosis-defense 与 proteostasis 轴，并标注 SLC7A11-GPX4-GSH 和 HSP90/UPR；BeatAML ex vivo pharmacogenomic associations。使用简洁 AML 细胞图标、转录组图标、网络图标和药敏板图标。背景白色，主色为蓝绿灰，少量红色或金色强调预后和药敏关联。右下角或底部必须加入边界框："Retrospective public datasets; ex vivo association; experimental validation required; no clinical treatment recommendation." 不要画真实患者照片，不要画过度复杂通路，不要使用 proved mechanism、causes、activates、clinical recommendation 等表达。

# Layout notes
- Suggested aspect ratio: horizontal 16:9 or journal graphical-abstract rectangle.
- Keep labels short; avoid overcrowding.
- Use "state model" rather than "mechanistic pathway".
- Place the six-gene signature in the central visual anchor.
"""


def cover_letter() -> str:
    return f"""
Dear Editor-in-Chief / Editorial Office,

We are pleased to submit our manuscript entitled "{TITLE}" for consideration as a Research article in Human Genomics.

Acute myeloid leukemia remains difficult to stratify because genomic risk, transcriptional state, immune context and treatment persistence do not always align. In this study, we asked whether a proteostasis-associated ferroptosis-tolerance (PRFT) state could define a reproducible adverse AML transcriptional program with prognostic, clinical and pharmacogenomic relevance.

This manuscript is not intended as another generic prognostic-signature report. We first define and audit a PRFT state from AML bulk transcriptomes, then derive and retain a fixed six-gene signature comprising {SIX_GENES}. The signature was evaluated in TCGA-LAML, GSE37642, GSE12417 and a combined GPL570 cohort, remained informative after clinical adjustment for age, sex and WBC_log10, and was supported by a multi-algorithm survival machine-learning framework. Additional interpretation layers link PRFT-high AML to monocyte-like/myeloid immune-suppressive transcriptional programs, proteostasis and ferroptosis-defense axes, network-level prioritization and BeatAML ex vivo pharmacogenomic associations.

We believe the work is suitable for Human Genomics because it integrates transcriptomics, machine learning, clinical covariate auditing, processed single-cell table-based evidence and pharmacogenomic association within a genomics-focused AML framework. We have deliberately kept the claims bounded: the analyses are retrospective and based on public datasets; BeatAML findings are ex vivo associations rather than clinical treatment recommendations; and experimental validation is required before mechanistic or clinical-response conclusions can be drawn.

The manuscript has not been published and is not under consideration elsewhere. All authors have approved the manuscript for submission. [AUTHOR CONFIRMATION NEEDED: confirm final author approval.] The authors declare no competing interests. [AUTHOR CONFIRMATION NEEDED: confirm competing interests.] Funding, author contributions, data repository information and code repository information are provided in the manuscript declarations and require final author confirmation before submission.

Thank you for considering our manuscript. We would be grateful for the opportunity to have it evaluated for publication in Human Genomics.

Sincerely,

[AUTHOR CONFIRMATION NEEDED: Corresponding author name]
[AUTHOR CONFIRMATION NEEDED: Corresponding author affiliation]
[AUTHOR CONFIRMATION NEEDED: Corresponding author email]
"""


def data_availability() -> str:
    return """
# Availability of data and materials

The study used publicly available, de-identified datasets and derived analysis outputs. TCGA-LAML transcriptomic and clinical data were obtained from public TCGA/GDC or UCSC Xena resources. GEO datasets GSE37642 and GSE12417 were used for formal external validation. GSE6891 and GSE14468 were audited but not used as formal validation cohorts because the locally available files lacked sufficient six-gene coverage and usable overall-survival endpoints. BeatAML expression and ex vivo drug-response data were used for pharmacogenomic association analyses. Processed single-cell tables from GSE116256 were used for processed single-cell table-based patient-level analyses. STRING was used for protein-protein interaction evidence.

The derived datasets supporting the conclusions of this article are included within the article and its supplementary tables, and will also be deposited in [AUTHOR CONFIRMATION NEEDED: repository name and URL] before publication. These derived files include PRFT scores, differential-expression and WGCNA outputs, candidate-gene lists, survival-validation summaries, clinical Cox outputs, machine-learning performance tables, PPI/network-prioritization tables, BeatAML pharmacogenomic association summaries, processed single-cell patient-level summaries, bulk signature/pathway summaries, figure source audits and excluded-analysis audit tables.

No newly generated raw human sequencing data were produced in this study. No individual participant-level identifiable data are included in the manuscript or supplementary files.

# Repository and citation actions
- Confirm the final public repository for derived tables and scripts: [AUTHOR CONFIRMATION NEEDED: GitHub/Zenodo/OSF/institutional repository].
- Add a persistent identifier or DOI before publication if a repository record is created: [AUTHOR CONFIRMATION NEEDED: DOI or stable URL].
- Ensure public source datasets are cited in the final reference list where required by the journal.

# Chinese author check
- 请确认最终用于存放 derived tables/source code 的仓库 URL 和 DOI。
- 不要填写不存在的 Zenodo/GitHub/OSF 地址；如果尚未上传，请保留占位符直到上传完成。
- 原始 TCGA/GEO/BeatAML/GSE116256/STRING 数据为公共来源，本研究不产生新的可识别个人数据。
"""


def code_availability() -> str:
    return """
# Code availability

Code generated for this study will be deposited in [AUTHOR CONFIRMATION NEEDED: GitHub/Zenodo/OSF repository] before publication. The repository will include scripts or reproducible notebooks for data preprocessing, PRFT scoring, differential-expression analysis, WGCNA reproduction, survival modeling, clinical covariate auditing, machine-learning model evaluation, PPI/network prioritization, BeatAML pharmacogenomic association analysis, processed single-cell table-based analyses, bulk signature/pathway scoring, figure assembly and submission-package audits.

The archived version, license and permanent identifier are currently pending: [AUTHOR CONFIRMATION NEEDED: archived version DOI, software license and release date]. No non-public raw clinical data or restricted individual-level data should be included in the code repository.
"""


def reproducibility_readme() -> str:
    return """
# AML PRFT reproducibility README draft

## Project overview
This repository contains the derived outputs and code required to reproduce the manuscript "A proteostasis-associated ferroptosis-tolerance signature defines adverse prognosis, myeloid immune suppression and ex vivo pharmacogenomic vulnerability in acute myeloid leukemia." The project defines a PRFT state, retains a fixed six-gene AML signature and audits prognostic, clinical, machine-learning, network, bulk signature, processed single-cell and BeatAML ex vivo pharmacogenomic evidence.

## Data sources
- TCGA-LAML transcriptomic and clinical data from public TCGA/GDC or UCSC Xena resources.
- GEO GPL570 validation datasets: GSE37642 and GSE12417.
- GSE6891 and GSE14468: audit-exclusion datasets only.
- BeatAML expression and ex vivo drug-response data.
- Processed single-cell tables from GSE116256.
- STRING protein-protein interaction data.

## Directory structure
- 00_raw_data: original or downloaded data files when available locally.
- 01_processed_data: harmonized expression, clinical and derived intermediate data.
- 02_scripts: reproducible scripts by phase.
- 03_results_tables: tabular outputs and audit tables.
- 04_figures: generated figures and final main-figure composites.
- 05_logs: phase logs, checklists and limitation statements.
- 06_manuscript_support: manuscript text, figure legends and submission statements.
- 07_submission_package: final submission-package draft structure.

## Analysis phases
- Phase 0: project audit and file inventory.
- Phase 1/1-fix: PRFT scoring, DEG, WGCNA, 715 candidates, 33 cross-platform candidates and fixed six-gene model reproduction.
- Phase 2/2b: clinical independence and figure repair.
- Phase 3A/3B/3C: survival machine learning, PRFT-state recognition and BeatAML pharmacogenomic association.
- Phase 4/4b: processed single-cell table-based validation and patient-level robustness.
- Phase 5: PPI/network prioritization and FS12 feature strategy.
- Phase 6: GSE6891/GSE14468 audit exclusion.
- Phase 7: bulk signature-based immune/pathway inference.
- Phase 8A: alternative-splicing input audit and AS exclusion.
- Phase 9/10: manuscript integration, reference enhancement, figure assembly and submission text package.

## Reproducing main figures
Run the Python figure-assembly script after all phase outputs are present:

python 02_scripts/phase10D_figure_assembly.py

This generates Figure 1-7 composites under 04_figures/final_main_figures and copies them into 07_submission_package/figures. The current composites use existing source panels and table-derived schematic panels. They do not rerun Phase 1-9 analyses.

## Software versions
- R version: [AUTHOR CONFIRMATION NEEDED: final R version from sessionInfo].
- Python version: [AUTHOR CONFIRMATION NEEDED: final Python version].
- Key R/Python package versions: [AUTHOR CONFIRMATION NEEDED: final package version table or sessionInfo files].

## Limitations and excluded analyses
- The fixed six-gene signature is not replaced by later machine-learning or PPI analyses.
- BeatAML findings are ex vivo pharmacogenomic associations and should not be interpreted as clinical treatment recommendations.
- Processed single-cell analyses are table-based and do not constitute de novo single-cell reanalysis.
- Bulk immune/pathway analyses are signature-based inference and not formal immune deconvolution.
- Alternative splicing was excluded because no matched PSI/event-level input was available in the current analysis package.
- GSE6891 and GSE14468 were excluded from formal external validation after audit.

## Repository status
[AUTHOR CONFIRMATION NEEDED: final public repository URL, DOI, release date and license.]
"""


def declarations() -> str:
    return """
# Ethics approval and consent to participate

This study used publicly available, de-identified datasets and involved no new human participants or animal experiments. [AUTHOR CONFIRMATION NEEDED: institutional confirmation that additional ethics approval was not required.]

# Consent for publication

Not applicable.

# Competing interests

The authors declare no competing interests. [AUTHOR CONFIRMATION NEEDED]

# Funding

[AUTHOR CONFIRMATION NEEDED: list all funding sources, grant numbers and funding recipients. If no specific funding was received, confirm the appropriate statement.]

# Author contributions

Conceptualization: [AUTHOR CONFIRMATION NEEDED]

Data curation: [AUTHOR CONFIRMATION NEEDED]

Formal analysis: [AUTHOR CONFIRMATION NEEDED]

Methodology: [AUTHOR CONFIRMATION NEEDED]

Visualization: [AUTHOR CONFIRMATION NEEDED]

Writing - original draft: [AUTHOR CONFIRMATION NEEDED]

Writing - review & editing: [AUTHOR CONFIRMATION NEEDED]

Supervision: [AUTHOR CONFIRMATION NEEDED]

Funding acquisition: [AUTHOR CONFIRMATION NEEDED]

All authors read and approved the final manuscript. [AUTHOR CONFIRMATION NEEDED]

# Acknowledgements

[AUTHOR CONFIRMATION NEEDED: acknowledge colleagues, institutional support, language editing, data providers or computational resources as appropriate.]
"""


def submission_checklist() -> str:
    items = [
        ("Manuscript title page", "needs author metadata"),
        ("Structured abstract", "present in manuscript draft"),
        ("Keywords", "present in manuscript draft"),
        ("Main text", "reference-enhanced draft available"),
        ("References formatted", "61 references; final manual metadata check still recommended"),
        ("Figure legends", "Phase 10D legends generated"),
        ("Main Figures 1-7", "generated as PDF and PNG composites"),
        ("Supplementary Figures S1-S10", "plan and legends generated; final files require visual review"),
        ("Supplementary Tables S1-S15", "manifest ready; source files copied where available"),
        ("Cover letter", "generated in TXT and DOCX"),
        ("Graphical abstract", "design brief and prompt generated; final artwork requires visual approval"),
        ("Data availability", "draft generated; repository URL pending"),
        ("Code availability", "draft generated; repository URL and DOI pending"),
        ("Ethics statement", "template generated; institutional confirmation pending"),
        ("Author contributions", "CRediT template generated; author names pending"),
        ("Funding", "pending author confirmation"),
        ("Competing interests", "pending author confirmation"),
        ("Line numbers", "must be added in final manuscript formatting"),
        ("Page numbers", "must be added in final manuscript formatting"),
        ("Figure resolution", "Phase 10D PNG composites are 300 dpi; final visual QA required"),
        ("Figure file naming", "standardized under 07_submission_package/figures"),
        ("Supplementary file naming", "draft structure generated; final naming should follow submission system"),
        ("Final visual QA", "required before submission"),
        ("Repository URL", "pending author confirmation"),
        ("ORCID if required", "pending author confirmation"),
        ("Suggested reviewers if required", "optional; pending author decision"),
        ("Opposed reviewers if needed", "optional; pending author decision"),
    ]
    lines = ["# Human Genomics submission checklist", ""]
    for i, (item, status) in enumerate(items, 1):
        lines.append(f"{i}. {item}: {status}")
    lines.extend(
        [
            "",
            "# Journal-specific reminders",
            "- Human Genomics requests a cover letter explaining why the manuscript should be published in the journal, policy issues, competing interests, author approval and duplicate-submission status.",
            "- Main manuscript files should include line and page numbering before final submission.",
            "- Datasets supporting conclusions should be supplied through public repositories where appropriate, supplementary/additional files, or a clearly described availability route.",
        ]
    )
    return "\n".join(lines)


def author_confirmation_rows() -> list[dict[str, str]]:
    return [
        {
            "item": "Corresponding author name, affiliation and email",
            "required_for_submission": "yes",
            "current_status": "missing",
            "suggested_text_if_available": "[AUTHOR CONFIRMATION NEEDED]",
            "user_action_needed": "Provide final corresponding author details.",
            "priority": "must",
        },
        {
            "item": "All authors approved submission",
            "required_for_submission": "yes",
            "current_status": "placeholder in cover letter",
            "suggested_text_if_available": "All authors have approved the manuscript for submission.",
            "user_action_needed": "Confirm with all co-authors before submission.",
            "priority": "must",
        },
        {
            "item": "Competing interests",
            "required_for_submission": "yes",
            "current_status": "placeholder declaration",
            "suggested_text_if_available": "The authors declare no competing interests.",
            "user_action_needed": "Confirm whether any competing interests exist.",
            "priority": "must",
        },
        {
            "item": "Funding",
            "required_for_submission": "yes",
            "current_status": "missing",
            "suggested_text_if_available": "[AUTHOR CONFIRMATION NEEDED: grant names and numbers]",
            "user_action_needed": "Provide funding sources or confirm no specific funding.",
            "priority": "must",
        },
        {
            "item": "Author contributions",
            "required_for_submission": "yes",
            "current_status": "CRediT template generated",
            "suggested_text_if_available": "Conceptualization; Data curation; Formal analysis; Methodology; Visualization; Writing; Supervision; Funding acquisition.",
            "user_action_needed": "Assign author names to each CRediT role.",
            "priority": "must",
        },
        {
            "item": "Ethics approval confirmation",
            "required_for_submission": "yes",
            "current_status": "template generated",
            "suggested_text_if_available": "Publicly available de-identified datasets; no new human participants or animal experiments.",
            "user_action_needed": "Confirm local institutional wording.",
            "priority": "must",
        },
        {
            "item": "Data repository URL/DOI",
            "required_for_submission": "yes",
            "current_status": "missing",
            "suggested_text_if_available": "[AUTHOR CONFIRMATION NEEDED: repository URL]",
            "user_action_needed": "Create or confirm repository for derived tables and source data package.",
            "priority": "must",
        },
        {
            "item": "Code repository URL/DOI/license",
            "required_for_submission": "yes",
            "current_status": "missing",
            "suggested_text_if_available": "[AUTHOR CONFIRMATION NEEDED: GitHub/Zenodo/OSF]",
            "user_action_needed": "Create or confirm public code archive and license.",
            "priority": "must",
        },
        {
            "item": "ORCID identifiers",
            "required_for_submission": "no",
            "current_status": "not provided",
            "suggested_text_if_available": "[AUTHOR CONFIRMATION NEEDED]",
            "user_action_needed": "Provide ORCID IDs if required by the submission system.",
            "priority": "recommended",
        },
        {
            "item": "Suggested reviewers",
            "required_for_submission": "no",
            "current_status": "not provided",
            "suggested_text_if_available": "[OPTIONAL]",
            "user_action_needed": "Provide institutional emails/ORCID/Scopus IDs if suggesting reviewers.",
            "priority": "optional",
        },
        {
            "item": "Opposed reviewers",
            "required_for_submission": "no",
            "current_status": "not provided",
            "suggested_text_if_available": "[OPTIONAL]",
            "user_action_needed": "Provide names and reason only if necessary.",
            "priority": "optional",
        },
        {
            "item": "Final figure visual QA",
            "required_for_submission": "yes",
            "current_status": "Phase 10D composites generated",
            "suggested_text_if_available": "Manual visual approval required.",
            "user_action_needed": "Open Figure 1-7 PDFs/PNGs and approve final layout.",
            "priority": "must",
        },
        {
            "item": "Line and page numbering",
            "required_for_submission": "yes",
            "current_status": "pending final manuscript formatting",
            "suggested_text_if_available": "Add in final DOCX before submission.",
            "user_action_needed": "Apply final Human Genomics formatting.",
            "priority": "must",
        },
    ]


def save_author_confirmation_csv(rows: list[dict[str, str]]) -> Path:
    path = TABLE_DIR / "phase10E_author_confirmation_items.csv"
    with path.open("w", encoding="utf-8-sig", newline="") as fh:
        writer = csv.DictWriter(
            fh,
            fieldnames=[
                "item",
                "required_for_submission",
                "current_status",
                "suggested_text_if_available",
                "user_action_needed",
                "priority",
            ],
        )
        writer.writeheader()
        writer.writerows(rows)
    return path


def create_cover_letter_docx(text: str) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Cover Letter")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Human Genomics submission draft")
    r.font.name = "Arial"
    r.font.size = Pt(10)

    for block in text.strip().split("\n\n"):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.15
        run = para.add_run(block.replace("\n", " "))
        run.font.name = "Arial"
        run.font.size = Pt(11)

    path = MS_DIR / "phase10E_cover_letter_Human_Genomics.docx"
    doc.save(path)
    return path


def key_result_checklist() -> str:
    return """
1. Graphical abstract design brief generated: yes
2. Graphical abstract prompt generated: yes
3. Cover letter generated: yes
4. Data availability generated: yes
5. Code availability generated: yes
6. README generated: yes
7. Declarations template generated: yes
8. Human Genomics checklist generated: yes
9. Author confirmation table generated: yes
10. Fixed six-gene model preserved: yes
11. AS exclusion preserved: yes
12. Processed single-cell boundary preserved: yes
13. Bulk signature-based inference boundary preserved: yes
14. BeatAML ex vivo association boundary preserved: yes
15. Clinical recommendation avoided: yes
16. Recommend entering final manuscript formatting: yes, after author metadata and repository fields are confirmed
17. Recommend preparing submission: yes, after final visual QA and author confirmations
18. Issues still requiring human confirmation: corresponding author details; author order and contributions; funding; competing interests; ethics/institutional wording; data/code repository URL/DOI/license; ORCID; optional reviewer suggestions; final figure visual approval
19. Overall conclusion: Phase 10E generated the Human Genomics submission text package without adding analyses or changing the fixed PRFT model.
"""


def main() -> None:
    brief = write_text("phase10E_graphical_abstract_design_brief.txt", graphical_abstract_design_brief())
    prompt = write_text("phase10E_graphical_abstract_prompt_EN_CN.txt", graphical_abstract_prompt())
    cover_txt = write_text("phase10E_cover_letter_Human_Genomics.txt", cover_letter())
    cover_docx = create_cover_letter_docx(cover_letter())
    data = write_text("phase10E_data_availability_statement.txt", data_availability())
    code = write_text("phase10E_code_availability_statement.txt", code_availability())
    readme = write_text("phase10E_reproducibility_README.txt", reproducibility_readme())
    decl = write_text("phase10E_declarations_template.txt", declarations())
    checklist = write_text("phase10E_Human_Genomics_submission_checklist.txt", submission_checklist())
    author_csv = save_author_confirmation_csv(author_confirmation_rows())
    phase_log = LOG_DIR / "phase10E_key_result_checklist.txt"
    phase_log.write_text(key_result_checklist().strip() + "\n", encoding="utf-8")

    copy_if_exists(brief, "graphical_abstract")
    copy_if_exists(prompt, "graphical_abstract")
    copy_if_exists(cover_txt, "cover_letter")
    copy_if_exists(cover_docx, "cover_letter")
    copy_if_exists(data, "data_code_availability")
    copy_if_exists(code, "data_code_availability")
    copy_if_exists(readme, "data_code_availability")
    copy_if_exists(decl, "manuscript")
    copy_if_exists(checklist, "reporting_checklists")
    copy_if_exists(author_csv, "reporting_checklists")
    copy_if_exists(phase_log, "reporting_checklists")


if __name__ == "__main__":
    main()
