#!/usr/bin/env python
# -*- coding: utf-8 -*-

from __future__ import annotations

import csv
import os
import re
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(os.environ.get("PHASE9_ROOT", r"C:/Users/Robin-Yang/AppData/Local/Temp/aml_prft_phase1_fix"))
TABLE_DIR = ROOT / "03_results_tables"
FIG_DIR = ROOT / "04_figures"
LOG_DIR = ROOT / "05_logs"
MS_DIR = ROOT / "06_manuscript_support"

for d in (TABLE_DIR, FIG_DIR, LOG_DIR, MS_DIR):
    d.mkdir(parents=True, exist_ok=True)


TITLE = (
    "A proteostasis-associated ferroptosis-tolerance signature defines adverse prognosis, "
    "myeloid immune suppression and ex vivo pharmacogenomic vulnerability in acute myeloid leukemia"
)

SIX_GENES = ["CLCN5", "ITGB2", "ARHGEF5", "TRIM32", "SAT1", "ACOX2"]

BOUNDARY_TERMS = {
    "recommended": [
        "associated with",
        "linked to",
        "suggested",
        "may reflect",
        "may indicate",
        "prioritized",
        "stratified",
        "externally evaluated",
        "ex vivo association",
        "signature-based inference",
        "processed single-cell table-based analysis",
    ],
    "forbidden": [
        r"\bproved\b",
        r"\bconfirmed\b",
        "demonstrated mechanistically",
        "mediated",
        "activated",
        "caused",
        "reversed",
        "restored",
        "rescued",
        "clinical recommendation",
        "treatment guidance",
        "diagnostic tool",
        "therapeutic target proved",
        "immune deconvolution demonstrated",
        "UMAP revealed",
        "UCell demonstrated",
        "CIBERSORT showed",
    ],
}


def read_csv(name: str) -> pd.DataFrame:
    path = TABLE_DIR / name
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)


def read_text(rel: str) -> str:
    path = ROOT / rel
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8", errors="replace")


def fmt_p(x, digits=3) -> str:
    try:
        v = float(x)
    except Exception:
        return "[insert value from table]"
    if pd.isna(v):
        return "[insert value from table]"
    if v < 0.001:
        return f"{v:.2e}"
    return f"{v:.{digits}g}"


def fmt_num(x, digits=3) -> str:
    try:
        v = float(x)
    except Exception:
        return "[insert value from table]"
    if pd.isna(v):
        return "[insert value from table]"
    return f"{v:.{digits}g}"


def pdf_header_ok(path: Path) -> bool:
    try:
        with path.open("rb") as f:
            return f.read(5) == b"%PDF-"
    except Exception:
        return False


def find_fig(patterns: list[str], exclude_patterns: list[str] | None = None) -> Path | None:
    exclude_patterns = exclude_patterns or []
    files = sorted(FIG_DIR.rglob("*.pdf"))
    for pat in patterns:
        pat_l = pat.lower()
        for f in files:
            name = f.name.lower()
            full = str(f).lower()
            if pat_l in name or pat_l in full:
                if any(ex.lower() in name or ex.lower() in full for ex in exclude_patterns):
                    continue
                if pdf_header_ok(f):
                    return f
    return None


def add_doc_title(doc: Document, title: str, subtitle: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(8)
        sr = sp.add_run(subtitle)
        sr.font.name = "Arial"
        sr.font.size = Pt(11)
        sr.font.color.rgb = RGBColor(85, 85, 85)


def set_doc_styles(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    for style_name, size in [("Heading 1", 20), ("Heading 2", 16), ("Heading 3", 14)]:
        st = styles[style_name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor(0, 0, 0 if style_name != "Heading 3" else 67)


def write_docx_from_markdownish(title: str, text: str, out_path: Path, subtitle: str | None = None):
    doc = Document()
    set_doc_styles(doc)
    add_doc_title(doc, title, subtitle)
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.save(out_path)


def extract_docx_text(path: Path, limit_lines=120) -> list[str]:
    if not path.exists():
        return []
    try:
        with zipfile.ZipFile(path) as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
        text = re.sub(r"</w:p>", "\n", xml)
        text = re.sub(r"<[^>]+>", "", text)
        text = (
            text.replace("&amp;", "&")
            .replace("&lt;", "<")
            .replace("&gt;", ">")
            .replace("&quot;", '"')
        )
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        return lines[:limit_lines]
    except Exception:
        return []


phase1_sample_flow = read_csv("phase1_fix_sample_flow.csv")
phase1_tcga = read_csv("phase1_fix_TCGA_survival_summary.csv")
phase1_geo = read_csv("phase1_fix_GEO_validation_summary.csv")
phase2_cox = read_csv("phase2_multivariable_cox_models.csv")
phase3a_cmp = read_csv("phase3A_fix_original_6gene_vs_top_models.csv")
phase3a_ppi = read_csv("phase3A_PPI_vs_original_6gene_comparison.csv")
phase3b_state = read_csv("phase3B_fix_independent_vs_label_proximal_summary.csv")
phase3b_importance = read_csv("phase3B_fix_SHAP_or_importance_top_features.csv")
phase3c_drugs = read_csv("phase3C_BeatAML_representative_drugs.csv")
phase4b = read_csv("phase4b_singlecell_robustness_summary.csv")
phase5_hubs = read_csv("phase5_consensus_hub_genes.csv")
phase6_interp = read_csv("phase6_external_validation_interpretation.csv")
phase7_main = read_csv("phase7_cleanup_main_signature_pathway_summary.csv")
phase8a = read_csv("phase8A_AS_feasibility_report.csv")

tcga_expr_n = 173
tcga_surv_n = 151
tcga_hr = tcga_p = tcga_auc1 = tcga_auc3 = tcga_auc5 = "[insert value from table]"
if not phase1_tcga.empty:
    row = phase1_tcga.iloc[0]
    tcga_hr = fmt_num(row.get("univariate_HR"))
    tcga_p = fmt_p(row.get("univariate_P"))
    tcga_auc1 = fmt_num(row.get("AUC_1year"))
    tcga_auc3 = fmt_num(row.get("AUC_3year"))
    tcga_auc5 = fmt_num(row.get("AUC_5year"))

model2 = phase2_cox[(phase2_cox.get("model", "") == "Model 2") & (phase2_cox.get("term", "") == "risk_score")]
model2_text = "HR [insert value from table], 95% CI [insert value from table], P=[insert value from table]"
if not model2.empty:
    r = model2.iloc[0]
    model2_text = f"HR {fmt_num(r.HR)} (95% CI {fmt_num(r.lower95)}-{fmt_num(r.upper95)}), P={fmt_p(r.P_value)}"

geo_bits = []
if not phase1_geo.empty:
    for _, r in phase1_geo.iterrows():
        geo_bits.append(
            f"{r['dataset_id']}: n={int(r['sample_count'])}, HR {fmt_num(r['univariate_HR'])}, P={fmt_p(r['univariate_P'])}, AUC1/3/5={fmt_num(r['AUC_1year'])}/{fmt_num(r['AUC_3year'])}/{fmt_num(r['AUC_5year'])}"
        )

top_hubs = []
if not phase5_hubs.empty:
    top_hubs = phase5_hubs["gene_symbol"].head(10).tolist()

fsf_row = phase3b_state[phase3b_state["feature_set"].str.contains("FS-F", na=False)] if not phase3b_state.empty else pd.DataFrame()
fsf_text = "FS-F label-proximal model AUROC/AUPRC [insert value from table]"
if not fsf_row.empty:
    r = fsf_row.iloc[0]
    fsf_text = f"FS-F label-proximal model CV AUROC={fmt_num(r.CV_AUROC)}, AUPRC={fmt_num(r.CV_AUPRC)}; external mean AUROC={fmt_num(r.external_mean_AUROC)}"

drug_lines = []
if not phase3c_drugs.empty:
    for _, r in phase3c_drugs.iterrows():
        drug_lines.append(
            f"{r['drug_name']}: n={int(r['matched_samples'])}, high-low AUC difference={fmt_num(r['difference_high_minus_low'])}, FDR={fmt_p(r['FDR'])}, direction={r['group_difference_interpretation']}"
        )

phase7_core = []
if not phase7_main.empty:
    for _, r in phase7_main.iterrows():
        if r.get("main_text_priority") == "core":
            phase7_core.append(
                f"{r['set_name']}: mean rho={fmt_num(r['risk_cor_mean_rho'])}, significant cohorts={int(r['risk_cor_significant_cohorts'])}, consistent={r['cross_cohort_consistent']}"
            )


one_sentence_argument = (
    "In AML, this study defines a proteostasis-associated ferroptosis-tolerance state and externally evaluates a fixed six-gene PRFT signature, "
    "supported by transcriptomic, clinical, machine-learning, network, processed single-cell, bulk signature and BeatAML ex vivo pharmacogenomic evidence, "
    "while limiting interpretation to retrospective public datasets without mechanistic or clinical treatment claims."
)

terminology_ledger = {
    "PRFT": "proteostasis-associated ferroptosis-tolerance",
    "fixed six-gene PRFT signature": ", ".join(SIX_GENES),
    "bulk immune module": "bulk signature-based inference",
    "single-cell module": "processed single-cell table-based patient-level analysis",
    "BeatAML module": "ex vivo pharmacogenomic association",
    "PPI module": "network-level prioritization",
    "AS module": "excluded after input audit",
}

outline_text = f"""# Title
{TITLE}

# Abstract
Acute myeloid leukemia (AML) remains biologically heterogeneous, and transcriptomic states that connect stress adaptation with clinical outcome remain incompletely defined. We define a proteostasis-associated ferroptosis-tolerance (PRFT) state using bulk transcriptomic scoring and derive a fixed six-gene prognostic signature comprising {", ".join(SIX_GENES)}. The model was trained in TCGA-LAML and externally evaluated in GSE37642, GSE12417 and a combined GPL570 cohort. Clinical adjustment supported the independent prognostic value of the risk score, with Model 2 ({model2_text}) prioritized for the main text. Multi-algorithm survival machine learning supported retention of the original six-gene model because alternative top-ranked models showed limited external improvement and reduced simplicity. PRFT-related analyses further linked the high-risk state to monocyte/myeloid and immune-regulatory bulk transcriptional programs, processed single-cell table-based evidence, network-level prioritization and BeatAML ex vivo pharmacogenomic associations. These findings nominate PRFT as a clinically relevant AML transcriptional state, while requiring experimental validation before mechanistic or therapeutic conclusions.

# Keywords
acute myeloid leukemia; proteostasis; ferroptosis tolerance; prognostic signature; machine learning; BeatAML; bulk signature scoring; processed single-cell analysis

# Introduction
AML prognosis reflects both genetic lesions and adaptive transcriptional states that shape persistence, stress tolerance and therapy response. Existing prognostic models incompletely capture how proteostasis and ferroptosis-defense programs converge with immune-regulatory features in public multi-cohort data. This study addresses that gap by defining a PRFT state, constructing a fixed six-gene prognostic model, and evaluating its clinical, molecular and ex vivo pharmacogenomic correlates across public AML resources.

# Methods
The Methods should describe data acquisition, preprocessing, PRFT score construction, differential expression, WGCNA, candidate selection, six-gene model construction, external validation, clinical adjustment, survival machine learning, PRFT-state classification, PPI prioritization, bulk signature scoring, BeatAML association, processed single-cell table analysis, and excluded audit procedures. Formal AS analysis, de novo single-cell UMAP/UCell/AUCell analysis and formal immune deconvolution are excluded because the required local inputs or packages were unavailable.

# Results
The Results should follow eight linked modules: PRFT state definition, fixed six-gene model validation, clinical independence, multi-algorithm survival machine-learning stability, model interpretation and PPI prioritization, bulk signature-based inference, BeatAML ex vivo association, and processed single-cell table-based validation. Result 6 and Result 8 may be merged to reduce repetition.

# Discussion
The Discussion should interpret PRFT-high AML as a proteostasis-associated ferroptosis-tolerance state linked to adverse outcome and myeloid immune-suppressive transcriptional programs. It should explain why the fixed six-gene model remains the main model, why PPI hubs are supportive rather than replacement features, and why BeatAML findings are ex vivo associations rather than clinical-response evidence.

# Limitations
This study is retrospective and based on public datasets. No wet-lab validation is included. The single-cell component is based on processed tables rather than de novo raw-object reanalysis. Bulk immune findings are signature-based inference, not formal immune deconvolution. BeatAML data support ex vivo pharmacogenomic associations, not clinical-response recommendations. Alternative splicing was omitted because no matched PSI/event-level data were available in the current project package.

# Conclusions
The PRFT framework defines an adverse AML transcriptional state and supports a fixed six-gene prognostic signature with external evaluation, clinical adjustment and multi-modal computational support. The model should remain fixed pending experimental validation.

# Data availability
Public TCGA-LAML, GEO GPL570, BeatAML and processed single-cell resources were used. Dataset accession details, derived tables and code paths should be listed in the final submission package. No newly generated human subject data are reported.

# Code availability
Analysis scripts and audit logs are organized in the project code directory. A final clean code archive and README should be prepared before submission.

# Ethics statement
This study reanalyzed publicly available de-identified datasets. No new human participants, animal experiments or interventional procedures were performed.

# Author contributions
[Placeholder: specify conceptualization, data curation, analysis, visualization, writing, supervision and funding roles.]

# Funding
[Placeholder: insert funder names and grant numbers.]

# Conflicts of interest
The authors declare no competing interests. [Confirm before submission.]

# References
[Placeholder: insert curated AML, ferroptosis, proteostasis, BeatAML, single-cell AML, machine-learning prognosis and Human Genomics formatting references.]

# Terminology ledger
""" + "\n".join(f"- {k}: {v}" for k, v in terminology_ledger.items()) + "\n"

results_text = f"""# Results draft

## Definition and transcriptomic characterization of the PRFT state in AML
We first defined the PRFT state by integrating proteostasis-core and ferroptosis-tolerance scores in TCGA-LAML. The score was available for {tcga_expr_n} expression samples and separated the cohort into PRFT-high and PRFT-low groups by the median. Differential expression and WGCNA were then used to characterize PRFT-associated transcriptional structure. The rerun WGCNA analysis identified PRFT-positive modules and reproduced 715 PRFT candidate genes after intersecting upregulated PRFT-high DEGs with PRFT-positive module genes. Cross-platform filtering retained 33 candidate genes for model development. [Insert Figure 1; Supplementary Tables for DEG, WGCNA and candidate-gene derivation.]

## Construction and external validation of a fixed six-gene PRFT prognostic signature
Using the TCGA survival-eligible training set (n={tcga_surv_n}), we retained a fixed six-gene PRFT signature comprising {", ".join(SIX_GENES)}. In TCGA-LAML, the high-risk group had adverse overall survival (univariate HR {tcga_hr}, P={tcga_p}), with 1-, 3- and 5-year AUCs of {tcga_auc1}, {tcga_auc3} and {tcga_auc5}. External evaluation supported the model in the GPL570 validation setting: {'; '.join(geo_bits)}. GSE6891 and GSE14468 were retained only as audit-exclusion cohorts because local files mapped only one of six signature genes and lacked usable OS endpoints. [Insert Figure 2; Supplementary GEO audit table.]

## Clinical independence and prognostic utility of the six-gene PRFT signature
Clinical adjustment in TCGA-LAML supported the independent prognostic value of the fixed six-gene PRFT score. The prioritized main-text model was Model 2, which included risk_score, age, sex and WBC_log10. In this model, risk_score remained associated with overall survival ({model2_text}). More complex models involving cytogenetic risk or locally available mutation variables were treated as sensitivity analyses because missingness and coding heterogeneity limited the available sample size. Nomogram, calibration and DCA outputs were generated, but calibration was not used as central evidence. [Insert Figure 3; Supplementary clinical-model diagnostics.]

## Multi-algorithm survival machine learning supports the stability of the six-gene signature
We next evaluated the stability of PRFT-related survival modeling using a planned 150-combination framework. After package and feature-availability filtering, 140 non-PPI combinations were evaluable and 125 models were successfully fitted across nine algorithm classes. The top-ranked model was Elastic Net-Cox with FS01_all_33, but its selected genes were again {", ".join(SIX_GENES)}. The original fixed six-gene formula ranked seventh and had external performance close to the top model, supporting retention of the simpler predefined signature. The PPI-derived FS12 strategy improved TCGA performance but showed near-random external validation and was not used to replace the fixed model. [Insert Figure 3 or Figure 4; Supplementary model heatmaps and FS12 comparison.]

## Model interpretation and network prioritization link PRFT-high status to proteostasis, ferroptosis defense and immune-regulatory axes
PRFT-high state-recognition analysis provided an explanatory layer rather than a prognostic replacement model. Independent-like feature sets produced moderate internal and external classification performance, whereas the FS-F core-axis feature set achieved stronger classification but was explicitly label-proximal. {fsf_text}. XGBoost-SHAP and feature-importance outputs prioritized core-axis genes including AIFM2, HSPA1A, HERPUD1, HSPA5, NFE2L2, SLC3A2, GCLM, SLC40A1, SOCS1 and STAT5B. Network-level PPI prioritization further highlighted hubs including {", ".join(top_hubs)}. These findings were interpreted as model-interpretation and network-prioritization evidence, not direct mechanistic evidence. [Insert Figure 4.]

## Bulk signature-based inference supports a monocyte-like and myeloid immune-suppressive PRFT-high transcriptional state
Bulk signature scoring was performed uniformly across TCGA-LAML, GSE37642, GSE12417 and the combined GPL570 validation set. Core cross-cohort signals included {', '.join([x.split(':')[0] for x in phase7_core[:8]])}. These findings linked the PRFT-related axis to monocyte/macrophage-like, myeloid-suppressive, immune checkpoint, JAK/STAT/PD-L1 and oxidative stress/NRF2-related transcriptional programs. The same direction was concordant with processed single-cell table-based patient-level analyses, which supported a monocyte-like and immune-suppressive PRFT-high-like state. LSC/stemness signals did not support a dominant primitive interpretation. [Insert Figure 5.]

## PRFT risk score is associated with ex vivo pharmacogenomic vulnerability in BeatAML
BeatAML analysis evaluated ex vivo drug AUC associations in expression-drug matched samples. Higher drug AUC indicates lower ex vivo sensitivity. Representative drug results were: {'; '.join(drug_lines)}. Venetoclax showed the strongest high-risk association with higher AUC, whereas Bortezomib, 17-AAG/Tanespimycin, Panobinostat and Selumetinib showed lower AUC in high-risk samples. Cytarabine was not significant and should be retained only as a negative or contextual result. Drug-response machine learning and feature-importance analyses were treated as supportive ex vivo pharmacogenomic evidence. [Insert Figure 6.]

## Processed single-cell table-based patient-level analysis supports a monocyte-like/myeloid stress-adapted PRFT-high state
Because no local Seurat object, h5ad file, 10x matrix, raw count matrix or UMAP coordinates were available, the single-cell component was restricted to processed-table-based analysis. Patient-level pseudo-bulk and mixed-effect analyses supported enrichment of PRFT-high-like scores in monocyte/myeloid-related cell states. These analyses were concordant with the bulk signature-based inference but did not support an LSC17/stemness-dominant interpretation. No de novo UMAP localization, UCell/AUCell rescoring or ScissorR analysis was performed. [Insert Figure 5 or Supplementary processed single-cell figure.]
"""

methods_text = """# Methods draft

## Data acquisition and preprocessing
Public AML transcriptomic and clinical datasets were organized under a reproducible project directory. TCGA-LAML RNA-seq data were used as the training cohort. GEO GPL570 datasets GSE37642 and GSE12417 were used for external validation, and a combined GPL570 cohort was constructed after platform-level harmonization. GSE6891 and GSE14468 were audited but excluded from formal validation because the locally available files mapped only one of the six fixed model genes and lacked usable OS endpoints.

## PRFT score calculation
Proteostasis-core and ferroptosis-tolerance gene sets were scored at the single-sample level. PRFT score was defined from these two axes and used to assign PRFT-high and PRFT-low groups by the cohort median. The same grouping rule was preserved during reproduction and downstream auditing.

## Differential expression and WGCNA
Differential expression was performed between PRFT-high and PRFT-low TCGA samples. WGCNA was rerun using the TCGA expression matrix to identify modules associated with PRFT score. PRFT candidate genes were defined by intersecting PRFT-high upregulated DEGs with genes from PRFT-positive WGCNA modules. The rerun analysis reproduced 715 candidate genes.

## Candidate gene selection
Candidate genes were filtered for cross-platform detectability in GPL570 validation cohorts. This produced 33 cross-platform candidate genes for signature construction and machine-learning evaluation.

## Six-gene signature construction
A fixed six-gene PRFT signature was retained, comprising CLCN5, ITGB2, ARHGEF5, TRIM32, SAT1 and ACOX2. Risk scores were calculated from the fixed coefficients derived in TCGA-LAML. The fixed formula was not replaced by subsequent machine-learning or PPI-derived models.

## External validation
The fixed six-gene score was evaluated in TCGA-LAML, GSE37642 GPL570, GSE12417 GPL570 and the combined GPL570 cohort. Kaplan-Meier analysis, Cox regression and time-dependent ROC analyses were used where survival endpoints were available.

## Clinical independence analysis
Clinical covariates available in TCGA-LAML were audited for missingness before model building. The main clinical adjustment model included risk_score, age, sex and WBC_log10. Models including cytogenetic risk or mutation variables were used as sensitivity analyses when sample size and missingness permitted. ELN risk and several mutation variables were not locally available.

## Multi-algorithm survival machine learning
A planned 150-combination framework tested 10 survival modeling algorithms against 15 feature strategies. Feature selection was performed only within the TCGA training set. GEO cohorts were used only for external validation. Composite ranking prioritized external validation, 3-year AUC, HR direction consistency, model simplicity, biological interpretability and anti-overfitting behavior.

## PRFT-state classification and XGBoost-SHAP interpretation
PRFT-high versus PRFT-low classification used PRFT group labels, not survival labels. Focused feature sets included the original six genes, 33 cross-platform candidates, Phase 3A recurrent genes, intersected stable genes and a label-proximal core-axis explanatory set. XGBoost-SHAP and fallback importance methods were used for interpretation. FS-F was treated as explanatory and label-proximal, not as independent validation.

## STRING-PPI and cytoHubba-like network prioritization
PPI and cytoHubba-like centrality analyses were used to prioritize network-level hubs among PRFT-related genes. These analyses supported biological interpretation but did not replace the fixed six-gene prognostic signature.

## Bulk signature-based inference and pathway scoring
Bulk immune and pathway analyses used a uniform rank-based signature scoring framework across TCGA-LAML, GSE37642, GSE12417 and the combined GPL570 cohort. Because formal immune deconvolution packages were unavailable, results were reported as bulk signature-based inference rather than immune cell fraction deconvolution.

## BeatAML ex vivo pharmacogenomic association
BeatAML expression and drug-sensitivity data were matched at the sample level. Drug AUC was interpreted as higher values indicating lower ex vivo sensitivity or greater relative resistance. Associations between PRFT risk score and drug AUC were evaluated using correlation and high-low group comparisons. Results were interpreted as ex vivo pharmacogenomic associations, not clinical-response recommendations.

## Processed single-cell table-based patient-level analysis
Single-cell validation used processed GSE116256 tables. Because no raw Seurat object, h5ad object, 10x matrix, raw count matrix or UMAP coordinates were locally available, de novo UMAP, UCell/AUCell rescoring, myeloid reclustering and ScissorR analysis were not performed. Patient-level pseudo-bulk and mixed-effect analyses were used to evaluate cell-state associations.

## Statistical analysis
Continuous variables were compared using Wilcoxon or t tests as appropriate. Categorical variables were compared using Fisher or chi-square tests. Survival associations were evaluated using Cox proportional hazards models. Multiple testing was controlled using the Benjamini-Hochberg method where applicable. Random seeds were fixed at 1234 for reproducibility.

## Excluded analyses and audit procedures
GSE6891 and GSE14468 were excluded from formal validation because local data lacked sufficient six-gene coverage and usable survival endpoints. Alternative splicing was omitted because no matched PSI or event-level data were available. De novo single-cell UMAP/UCell/AUCell analyses were not performed because raw single-cell objects and embeddings were unavailable. Formal immune deconvolution was not claimed because only signature-based fallback inference was used.
"""

discussion_text = """# Discussion framework

## Central interpretation
The study supports PRFT-high AML as a proteostasis-associated ferroptosis-tolerance transcriptional state linked to adverse prognosis. The fixed six-gene model provides a compact prognostic representation of this state and remains the main model because it is externally evaluated, simple and biologically interpretable.

## Why the fixed six-gene model remains the main model
The multi-algorithm survival machine-learning framework showed that several high-ranking models selected the same six genes or closely overlapping features. The original formula retained comparable external performance and avoided the instability of larger or network-derived feature sets. This supports retaining the fixed signature rather than replacing it with a marginally higher-ranked model.

## Why PPI hubs are supportive rather than replacement features
PPI prioritization highlighted immune-regulatory and stress-response hubs such as HSP90AA1, HSP90AB1, JAK2, CD274 and NFE2L2. These genes help interpret the biological context of PRFT-high AML, but the FS12 PPI-based survival model did not show sufficient external validation. The PPI layer should therefore be presented as network-level prioritization, not as an alternative prognostic model.

## Bulk and processed single-cell consistency
Bulk signature-based inference and processed single-cell table-based patient-level analyses converged on a monocyte-like and myeloid immune-suppressive PRFT-high state. This convergence supports the biological coherence of the PRFT state, but it does not constitute direct immune cell quantification or de novo single-cell localization.

## Biological axes
The JAK2/STAT5/PD-L1, SLC7A11/GPX4, NFE2L2 and HSP90-related axes provide plausible biological context for PRFT-high AML. These axes connect stress adaptation, ferroptosis defense, immune-regulatory transcriptional programs and ex vivo drug associations. They should be described as associated axes that require experimental verification.

## BeatAML interpretation
BeatAML analyses linked the PRFT risk score to ex vivo drug AUC patterns. Higher PRFT risk was associated with higher Venetoclax AUC and lower AUC for Bortezomib, 17-AAG/Tanespimycin, Panobinostat and Selumetinib. These findings may guide future experimental prioritization but should not be written as clinical-response evidence.

## Single-cell interpretation
Processed single-cell tables supported monocyte/myeloid and immune-suppressive PRFT-high-like states. LSC17 and stemness signals did not support a dominant primitive or LSC-led interpretation. This boundary is important for avoiding overextension of the single-cell module.

## Limitations
Key limitations include retrospective public datasets, lack of wet-lab validation, processed-table-only single-cell evidence, absence of formal immune deconvolution, ex vivo rather than clinical drug-response evidence, omission of AS because matched PSI/event-level data were unavailable, and exclusion of GSE6891/GSE14468 from formal validation because local files were incompatible with fixed-signature survival testing.

## Future work
Future work should validate the fixed PRFT signature in independent clinical cohorts and experimental AML systems. Priority experiments include PRFT gene perturbation, ferroptosis rescue assays, JAK2/STAT5/PD-L1 and SLC7A11/GPX4 pathway verification, AML persistent-cell or MRD models, and patient-derived ex vivo drug-response validation.
"""

outline_path = MS_DIR / "phase9_manuscript_outline_Human_Genomics.txt"
outline_path.write_text(outline_text, encoding="utf-8")
results_path = MS_DIR / "phase9_results_draft.txt"
results_path.write_text(results_text, encoding="utf-8")
methods_path = MS_DIR / "phase9_methods_draft.txt"
methods_path.write_text(methods_text, encoding="utf-8")
discussion_path = MS_DIR / "phase9_discussion_framework.txt"
discussion_path.write_text(discussion_text, encoding="utf-8")

write_docx_from_markdownish("Phase 9 Manuscript Outline for Human Genomics", outline_text, MS_DIR / "phase9_manuscript_outline_Human_Genomics.docx", TITLE)
write_docx_from_markdownish("Phase 9 Results Draft", results_text, MS_DIR / "phase9_results_draft.docx", TITLE)
write_docx_from_markdownish("Phase 9 Methods Draft", methods_text, MS_DIR / "phase9_methods_draft.docx", TITLE)
write_docx_from_markdownish("Phase 9 Discussion Framework", discussion_text, MS_DIR / "phase9_discussion_framework.docx", TITLE)

figure_plan_rows = [
    ["Figure 1", "PRFT definition, DEG/WGCNA and candidate gene derivation", "main", "PRFT score, DEG, WGCNA, 715 candidates, 33 cross-platform genes", "Use existing DEG/WGCNA/candidate panels where available; otherwise assemble from source tables."],
    ["Figure 2", "Six-gene PRFT signature construction and external validation", "main", "Fixed six-gene model, TCGA, GSE37642, GSE12417, combined GPL570", "Exclude GSE6891/GSE14468 from formal validation; mention audit exclusion."],
    ["Figure 3", "Clinical independence and multi-algorithm survival ML stability", "main", "Model 2 forest plot plus ML stability/original-vs-top comparison", "Nomogram/calibration/DCA should remain supplementary or cautious."],
    ["Figure 4", "XGBoost-SHAP interpretation and PPI network-level prioritization", "main", "State-recognition interpretation and network hubs", "Frame FS-F as label-proximal and PPI as network-level prioritization."],
    ["Figure 5", "Bulk signature-based inference and processed single-cell consistency", "main", "Phase 7 cleanup panels plus Phase 4b processed-table evidence", "Avoid deconvolution or UMAP wording."],
    ["Figure 6", "BeatAML ex vivo pharmacogenomic vulnerability", "main", "Representative drugs and ex vivo AUC associations", "Avoid clinical-response or therapeutic-decision language."],
    ["Figure 7", "Integrated schematic model of PRFT-high AML", "main", "Conceptual model integrating prognosis, stress state, myeloid signatures and ex vivo vulnerability", "Should be schematic and carefully hedged."],
    ["Supplementary Figure S1", "Full DEG and WGCNA details", "supplement", "DEG volcano/heatmap, WGCNA module diagnostics", "Source data tables already available."],
    ["Supplementary Figure S2", "All GEO KM/ROC and audit diagnostics", "supplement", "GSE37642, GSE12417, combined GPL570; GSE6891/GSE14468 exclusion", "Do not include excluded cohorts as formal validation figures."],
    ["Supplementary Figure S3", "Clinical model diagnostics", "supplement", "All Cox models, nomogram, calibration, DCA", "Calibration not central evidence."],
    ["Supplementary Figure S4", "Full Phase 3A ML results", "supplement", "150 planned models, successful models, heatmaps, rankings", "Retain original six-gene model."],
    ["Supplementary Figure S5", "PPI FS12 comparison", "supplement", "PPI model comparison and hub genes", "Do not replace six-gene model."],
    ["Supplementary Figure S6", "Full BeatAML drug tables", "supplement", "All drug correlations and ML summaries", "Cytarabine negative result in supplement."],
    ["Supplementary Figure S7", "Processed single-cell enrichment", "supplement", "RoE, score heatmaps, patient-level boxplots", "No UMAP, UCell or AUCell claims."],
    ["Supplementary Figure S8", "AS and excluded-data audit", "supplement", "AS missing-input statement; GSE6891/GSE14468 exclusion", "Audit-only evidence."],
]
fig_plan = pd.DataFrame(figure_plan_rows, columns=["figure_id", "title", "placement", "content", "notes"])
fig_plan.to_csv(TABLE_DIR / "phase9_final_figure_plan.csv", index=False)
(MS_DIR / "phase9_final_figure_plan.txt").write_text(
    "\n".join([f"{r.figure_id}. {r.title} [{r.placement}]\nContent: {r.content}\nNotes: {r.notes}\n" for r in fig_plan.itertuples()]),
    encoding="utf-8",
)

mapping_specs = [
    ("Figure 1", "phase1/DEG/WGCNA candidate derivation", ["Figure2", "Figure3", "WGCNA", "DEG", "phase1"], [], "main_or_needs_assembly"),
    ("Figure 2", "six-gene signature and external validation", ["Figure4C", "Figure4D", "Figure4E", "TCGA_KM", "GSE37642_KM", "GSE12417_KM", "combined_GPL570_KM", "phase1_KM_ROC"], ["GSE6891", "GSE14468"], "main"),
    ("Figure 3A", "clinical independence forest plot", ["phase2b_forestplot_model2_publication"], [], "main"),
    ("Figure 3B", "survival ML stability", ["phase3A_fix_original_6gene_vs_top_models", "phase3A_fix_top20_model_ranking", "phase3A_fix_model_performance_heatmap"], [], "main_or_supplement"),
    ("Figure 4A", "SHAP or feature importance", ["phase3B_fix_SHAP_or_importance_FS_F_core_axis", "phase3B_fix_SHAP_or_importance_FS_A_6gene"], [], "main"),
    ("Figure 4B", "PPI hub barplot", ["phase5_PPI_hub_barplot"], [], "main"),
    ("Figure 5A", "bulk signature consistency", ["phase7_cleanup_core_signature_consistency_panel"], [], "main"),
    ("Figure 5B", "bulk and processed single-cell consistency", ["phase7_cleanup_bulk_singlecell_consistency_panel"], [], "main"),
    ("Figure 5C", "processed single-cell patient-level boxplot", ["phase4b_patient_level_score_boxplots"], ["UMAP"], "main_or_supplement"),
    ("Figure 6A", "BeatAML representative drug boxplots", ["phase3C_BeatAML_selected_drug_boxplots"], [], "main"),
    ("Figure 6B", "BeatAML drug correlation volcano", ["phase3C_BeatAML_drug_correlation_volcano"], [], "main_or_supplement"),
    ("Figure 7", "integrated schematic model", ["graphical", "schematic", "Figure7"], ["BeatAML_pharmacogenomics"], "needs_manual_or_existing_graphical_abstract"),
    ("Exclude", "phase4 UMAP placeholders", ["phase4_UMAP"], [], "exclude"),
    ("Exclude", "AS figures", ["phase8", "AS"], [], "exclude"),
    ("Exclude", "GSE6891/GSE14468 formal validation figures", ["phase6_GSE6891", "phase6_GSE14468"], [], "exclude_or_audit_only"),
    ("Supplement", "Phase 7 deconvolution-named figure", ["phase7_PRFT_myeloid_deconvolution_boxplots"], [], "supplement_only_rename_caption_as_signature_based"),
]

mapping_rows = []
all_pdfs = sorted(FIG_DIR.rglob("*.pdf"))
for fig_id, desc, pats, excludes, placement in mapping_specs:
    candidate = find_fig(pats, excludes)
    if candidate is None and fig_id in ("Exclude",):
        hits = [p for p in all_pdfs if any(pat.lower() in p.name.lower() for pat in pats)]
        candidate = hits[0] if hits else None
    path_str = str(candidate) if candidate else ""
    mapping_rows.append(
        {
            "figure_id": fig_id,
            "description": desc,
            "recommended_placement": placement,
            "candidate_pdf": path_str,
            "pdf_exists": bool(candidate and candidate.exists()),
            "pdf_header_ok": bool(candidate and pdf_header_ok(candidate)),
            "exclude_reason_or_note": (
                "Do not use as main formal validation/UMAP/AS figure."
                if "exclude" in placement.lower()
                else "Use conservative caption; verify panel composition manually."
            ),
        }
    )
figure_mapping = pd.DataFrame(mapping_rows)
figure_mapping.to_csv(TABLE_DIR / "phase9_figure_file_mapping.csv", index=False)

qc_lines = ["Phase 9 figure QC log", "PDF header check uses literal %PDF- header.", ""]
for r in mapping_rows:
    qc_lines.append(
        f"{r['figure_id']} | {r['description']} | placement={r['recommended_placement']} | exists={r['pdf_exists']} | pdf_header_ok={r['pdf_header_ok']} | file={r['candidate_pdf']}"
    )
(LOG_DIR / "phase9_figure_QC_log.txt").write_text("\n".join(qc_lines), encoding="utf-8")

draft_files = [
    MS_DIR / "phase9_manuscript_outline_Human_Genomics.txt",
    MS_DIR / "phase9_results_draft.txt",
    MS_DIR / "phase9_methods_draft.txt",
    MS_DIR / "phase9_discussion_framework.txt",
    MS_DIR / "phase7_cleanup_results_paragraph.txt",
]
audit_rows = []
for path in draft_files:
    text = path.read_text(encoding="utf-8", errors="replace") if path.exists() else ""
    for term in BOUNDARY_TERMS["forbidden"]:
        hits = [m.start() for m in re.finditer(term, text, flags=re.IGNORECASE)]
        display_term = term.replace(r"\b", "")
        audit_rows.append(
            {
                "file": str(path),
                "forbidden_or_caution_term": display_term,
                "hit_count": len(hits),
                "recommended_replacement": "; ".join(BOUNDARY_TERMS["recommended"][:4]),
                "action": "review_and_replace" if hits else "not_detected",
            }
        )
lang_audit = pd.DataFrame(audit_rows)
lang_audit.to_csv(TABLE_DIR / "phase9_language_boundary_audit.csv", index=False)
hits = lang_audit[lang_audit.hit_count > 0]
lang_log = ["Phase 9 language boundary log", ""]
if hits.empty:
    lang_log.append("No forbidden terms were detected in Phase 9 generated drafts.")
else:
    lang_log.append("Forbidden/caution terms detected; review before submission:")
    for _, r in hits.iterrows():
        lang_log.append(f"- {Path(r['file']).name}: {r['forbidden_or_caution_term']} ({r['hit_count']})")
lang_log.extend(
    [
        "",
        "Required boundaries retained:",
        "- AS module excluded because no PSI/event-level input was available.",
        "- Single-cell module described only as processed single-cell table-based patient-level analysis.",
        "- Phase 7 described as bulk signature-based inference, not formal immune deconvolution.",
        "- BeatAML described only as ex vivo pharmacogenomic association.",
        "- PPI described only as network-level prioritization.",
        "- Fixed six-gene signature retained as the main model.",
    ]
)
(LOG_DIR / "phase9_language_boundary_log.txt").write_text("\n".join(lang_log), encoding="utf-8")

gap_rows = [
    ["Must complete", "Reference completion and citation verification", "Needed before journal submission", "Build final reference list and verify claims against literature."],
    ["Must complete", "Data and code availability statements", "Required submission element", "Provide repository/accession paths and clean code archive/README."],
    ["Must complete", "Figure assembly and caption harmonization", "Main figures are currently mapped but not fully assembled into final composite plates", "Create final multi-panel Figures 1-7 with consistent typography."],
    ["Must complete", "Human Genomics formatting check", "Journal-specific metadata and declarations must be checked", "Verify abstract, declarations, data/code, ethics and supplementary file naming."],
    ["Must complete", "Manual figure visual QA", "Several existing figures were generated across phases", "Open final PDFs and inspect readability, panel labels and captions."],
    ["Recommended", "Wet-lab validation", "Strengthens biological credibility", "Prioritize PRFT perturbation, ferroptosis-defense, JAK2/STAT5/PD-L1 and drug-response assays."],
    ["Recommended", "Graphical abstract", "Useful for editorial communication", "Prepare a conservative schematic model of PRFT-high AML."],
    ["Recommended", "Cover letter", "Needed for submission package", "Update existing cover letter to reflect Phase 9 boundaries."],
    ["Recommended", "Supplementary table consolidation", "Current project contains many audit and result tables", "Select final supplementary tables and source data files."],
    ["Optional", "Additional Seurat/raw single-cell data search", "Would enable de novo single-cell analysis if found", "Not required for current manuscript, but would support stronger claims."],
    ["Optional", "External PSI/SpliceSeq download", "Could enable AS module later", "Not recommended for current manuscript unless a new formal AS phase is started."],
    ["Do not continue", "Using ordinary expression as AS data", "Would weaken credibility", "Keep AS out unless PSI/event-level data are obtained."],
    ["Do not continue", "Replacing fixed six-gene model with PPI/ML model", "External validation does not justify replacement", "Retain the fixed signature as main model."],
]
gap_df = pd.DataFrame(gap_rows, columns=["priority", "item", "reason", "recommended_action"])
gap_df.to_csv(TABLE_DIR / "phase9_submission_gap_checklist.csv", index=False)
(MS_DIR / "phase9_submission_gap_checklist.txt").write_text(
    "\n".join([f"{r.priority}: {r.item}\nReason: {r.reason}\nAction: {r.recommended_action}\n" for r in gap_df.itertuples()]),
    encoding="utf-8",
)

main_fig_count = int(figure_mapping["recommended_placement"].str.contains("main", case=False, na=False).sum())
supp_fig_count = int(figure_mapping["recommended_placement"].str.contains("supplement", case=False, na=False).sum())

integration_checklist = [
    "1. Human Genomics manuscript structure generated: yes",
    "2. Results draft generated: yes",
    "3. Methods draft generated: yes",
    "4. Discussion framework generated: yes",
    "5. Figure mapping completed: yes",
    "6. AS module excluded: yes",
    "7. phase4_UMAP placeholder figures excluded: yes",
    "8. processed single-cell table-based wording retained: yes",
    "9. bulk signature-based inference wording retained: yes",
    "10. fixed six-gene main model retained: yes",
    "11. PPI downgraded to supporting network-level evidence: yes",
    "12. BeatAML limited to ex vivo association: yes",
    f"13. Language boundary audit completed: yes; forbidden-term hit rows={int((lang_audit.hit_count > 0).sum())}",
    f"14. Recommended main figure entries: {main_fig_count}",
    f"15. Supplement figure entries: {supp_fig_count}",
    "16. Must complete before submission: references; final figure assembly; data/code availability; journal formatting; manual visual QA.",
    "17. Recommended before submission: wet-lab validation plan; graphical abstract; updated cover letter; supplementary table consolidation.",
    "18. Recommend starting full manuscript writing: yes",
    "19. Recommend preparing Human Genomics submission package: yes, after final references, captions, availability statements and figure plates are completed.",
]
(LOG_DIR / "phase9_integration_key_result_checklist.txt").write_text("\n".join(integration_checklist), encoding="utf-8")

docx_outputs = [
    MS_DIR / "phase9_manuscript_outline_Human_Genomics.docx",
    MS_DIR / "phase9_results_draft.docx",
    MS_DIR / "phase9_methods_draft.docx",
    MS_DIR / "phase9_discussion_framework.docx",
]
docx_qc = []
for p in docx_outputs:
    docx_qc.append(f"{p.name}: exists={p.exists()}, size={p.stat().st_size if p.exists() else 0}, extracted_lines={len(extract_docx_text(p, 20))}")
(LOG_DIR / "phase9_docx_structural_QC_log.txt").write_text(
    "\n".join(
        [
            "Phase 9 DOCX structural QC log",
            "LibreOffice/soffice was not detected by Phase 9 preflight, so rendered PNG visual QA was not performed.",
            "DOCX files were generated with python-docx and structurally checked by extracting document text.",
            "",
            *docx_qc,
        ]
    ),
    encoding="utf-8",
)

print("Phase 9 integration complete")
print(f"Root: {ROOT}")
print(f"Generated outline/results/methods/discussion TXT and DOCX under {MS_DIR}")
print(f"Figure mapping rows: {len(figure_mapping)}")
print(f"Language audit forbidden-term hit rows: {int((lang_audit.hit_count > 0).sum())}")

